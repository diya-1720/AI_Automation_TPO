import os
import io
import json
import re
import traceback
from typing import Dict, Any, List, Optional
from PIL import Image
import docx
from pypdf import PdfReader
import google.generativeai as genai

def extract_text_from_docx_bytes(contents: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(contents))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def extract_text_from_pdf_bytes(contents: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(contents))
        full_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                full_text.append(t)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def prepare_pil_image(image_bytes: bytes) -> Optional[Image.Image]:
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode != 'RGB':
            img = img.convert('RGB')
        return img
    except Exception as e:
        print(f"Error opening image: {e}")
        return None

def extract_report_data_from_evidence(
    text_notes: Optional[str] = "",
    document_files: Optional[List[Dict[str, Any]]] = None,
    document_images: Optional[List[Dict[str, Any]]] = None,
    event_photos: Optional[List[Dict[str, Any]]] = None,
    feedback_graph: Optional[Dict[str, Any]] = None,
    feedback_notes: Optional[str] = ""
) -> Dict[str, Any]:
    """
    Extracts report information using Gemini multi-modal vision & text model.
    Enforces strict zero-hallucination policy.
    """
    combined_text_evidence = []
    if text_notes and text_notes.strip():
        combined_text_evidence.append(f"--- USER NOTES ---\n{text_notes.strip()}")

    if feedback_notes and feedback_notes.strip():
        combined_text_evidence.append(f"--- USER FEEDBACK NOTES ---\n{feedback_notes.strip()}")

    if document_files:
        for doc_item in document_files:
            fname = doc_item.get("filename", "document")
            content = doc_item.get("bytes", b"")
            if fname.lower().endswith(".docx"):
                extracted = extract_text_from_docx_bytes(content)
            elif fname.lower().endswith(".pdf"):
                extracted = extract_text_from_pdf_bytes(content)
            elif fname.lower().endswith(".txt"):
                extracted = content.decode("utf-8", errors="ignore")
            else:
                extracted = ""
            if extracted.strip():
                combined_text_evidence.append(f"--- DOCUMENT FILE ({fname}) ---\n{extracted.strip()}")

    prompt_parts = []
    
    system_instruction = """
You are an expert academic documentation AI for university/college event reports.
Your task is to analyze all provided evidence (text notes, official notices, brochures, certificates, photographs, and feedback charts) and extract factual information to generate an official institutional activity report.

CRITICAL INSTRUCTIONS & ZERO-HALLUCINATION POLICY:
1. Preserve Factual Accuracy: Extract ONLY facts supported by the provided evidence (dates, names, venues, participant counts, organizers, titles).
2. HARD RULE - NO FABRICATION:
   - If the evidence does NOT state a value for a field (e.g. SWOT weaknesses, resource person, feedback stats, outcome bullets), leave that field as an EMPTY STRING ("").
   - NEVER invent or guess dates, participant numbers, speaker names, venues, or feedback percentages.
   - Do NOT fill fields with generic placeholders like "N/A" or "Not provided" unless necessary; prefer empty string "" if unsupported.
3. Multi-Source Evidence Synergy: Combine details from notes, notices, brochures, and circulars into one accurate representation. If official notice/brochure conflicts with informal notes, prefer official details.
4. Rewrite Professionally: Where evidence IS provided, rewrite it in formal, polished academic tone suitable for college records. Bullet points should be clean and start with '• '.
5. Feedback Graph Analysis: If a Feedback Graph image is provided, visually inspect the chart/graph and write a concise 2-4 sentence factual interpretation under "feedback_interpretation". If no feedback graph or feedback numbers are provided, set "feedback_interpretation" to "".
6. Checklist Determination: Set checklist fields to "[✓]" if the evidence includes that document/proof type, or "[✗]" if absent/not provided.

EXPECTED JSON SCHEMA FORMAT:
Return ONLY a raw JSON object (no markdown, no code fences) with these exact keys:
{
  "activity_name": "Formal title of the activity or empty",
  "date_time": "Date and time string or empty",
  "department": "Department / Committee or empty",
  "venue": "Venue / Location or empty",
  "participants": "Number of participants or target audience string or empty",
  "nature_of_activity": "Nature of activity (Workshop/Guest Lecture/Seminar/etc) or empty",
  "mode_of_activity": "Offline / Online / Hybrid or empty",
  "activity_incharge": "Name/designation of activity in-charge or empty",
  "activity_coordinator": "Name/designation of coordinator or empty",
  "resource_person": "Name and affiliation of speaker/resource person or empty",
  "objectives": "Formatted bullet points of objectives or empty",
  "target_audience": "Target audience description or empty",
  "event_schedule": "Event schedule or session timeline if present or empty",
  "methodology": "Execution process description paragraph or empty",
  "students_selected": "Students selected/placed list or description if present or empty",
  "outcomes": "Formatted bullet points of outcomes or empty",
  "activity_summary": "Cohesive executive summary paragraph or empty",
  "strengths": "Strengths bullet points or empty",
  "weaknesses": "Weaknesses bullet points or empty",
  "opportunities": "Opportunities bullet points or empty",
  "threats": "Threats bullet points or empty",
  "feedback_summary": "Overall summary of feedback or empty",
  "feedback_interpretation": "Short factual analysis of uploaded feedback graph image or empty",
  "notice_brochure_tick": "[✓]" or "[✗]" or "",
  "attendance_list_tick": "[✓]" or "[✗]" or "",
  "photos_tick": "[✓]" or "[✗]" or "",
  "certificate_tick": "[✓]" or "[✗]" or "",
  "feedback_form_tick": "[✓]" or "[✗]" or "",
  "feedback_analysis_tick": "[✓]" or "[✗]" or "",
  "news_letter_data_tick": "[✓]" or "[✗]" or "",
  "media_news_details_tick": "[✓]" or "[✗]" or "",
  "co_po_mapping_tick": "[✓]" or "[✗]" or "",
  "any_other_tick": "[✓]" or "[✗]" or ""
}
"""

    prompt_parts.append(system_instruction)

    if combined_text_evidence:
        prompt_parts.append("EVIDENCE TEXT & DOCUMENTS:\n" + "\n\n".join(combined_text_evidence))
    else:
        prompt_parts.append("EVIDENCE TEXT & DOCUMENTS: None provided in text form.")

    # Attach document images (Notices, Brochures, Certificates)
    if document_images:
        for idx, img_item in enumerate(document_images):
            pil_img = prepare_pil_image(img_item.get("bytes", b""))
            if pil_img:
                prompt_parts.append(f"INFORMATIONAL DOCUMENT IMAGE #{idx + 1} ({img_item.get('filename', 'notice')}):")
                prompt_parts.append(pil_img)

    # Attach feedback graph if present
    if feedback_graph:
        fg_img = prepare_pil_image(feedback_graph.get("bytes", b""))
        if fg_img:
            prompt_parts.append("FEEDBACK GRAPH IMAGE FOR VISUAL ANALYSIS:")
            prompt_parts.append(fg_img)

    # Attach event photos for context
    if event_photos:
        for idx, ep_item in enumerate(event_photos[:3]): # max 3 for AI context
            ep_img = prepare_pil_image(ep_item.get("bytes", b""))
            if ep_img:
                prompt_parts.append(f"EVENT PHOTOGRAPH FOR CONTEXT #{idx + 1}:")
                prompt_parts.append(ep_img)

    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt_parts)
        response_text = response.text.strip() if response and response.text else ""
        
        # Clean up markdown code blocks if model wrapped output in ```json
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        elif response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        response_text = response_text.strip()

        data = json.loads(response_text)
        
        # Auto-set checklist ticks based on attached evidence if missing
        if document_images and not data.get("notice_brochure_tick"):
            data["notice_brochure_tick"] = "[✓]"
        if event_photos and not data.get("photos_tick"):
            data["photos_tick"] = "[✓]"
        if feedback_graph and not data.get("feedback_analysis_tick"):
            data["feedback_analysis_tick"] = "[✓]"

        return data

    except Exception as e:
        print(f"Error during AI extraction: {e}")
        traceback.print_exc()
        # Return fallback empty structure - NO fake data
        return {
            "activity_name": "",
            "date_time": "",
            "department": "",
            "venue": "",
            "participants": "",
            "nature_of_activity": "",
            "mode_of_activity": "",
            "activity_incharge": "",
            "activity_coordinator": "",
            "resource_person": "",
            "objectives": "",
            "target_audience": "",
            "event_schedule": "",
            "methodology": "",
            "students_selected": "",
            "outcomes": "",
            "activity_summary": "",
            "strengths": "",
            "weaknesses": "",
            "opportunities": "",
            "threats": "",
            "feedback_summary": "",
            "feedback_interpretation": "",
            "notice_brochure_tick": "[✓]" if document_images else "",
            "photos_tick": "[✓]" if event_photos else "",
            "feedback_analysis_tick": "[✓]" if feedback_graph else ""
        }
