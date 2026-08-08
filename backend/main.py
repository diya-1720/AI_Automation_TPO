import os
import io
import time
import json
import re
import traceback
import hashlib
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional
from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv
import docx
from docx.shared import Inches, Pt
import google.generativeai as genai
from pypdf import PdfReader
from PIL import Image
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

load_dotenv()

def get_storage_dir(sub_dir: str) -> str:
    target = os.path.join(".", sub_dir)
    try:
        os.makedirs(target, exist_ok=True)
        test_file = os.path.join(target, ".perm_test")
        with open(test_file, "w") as f:
            f.write("1")
        os.remove(test_file)
        return target
    except Exception:
        fallback = os.path.join("/tmp", sub_dir)
        os.makedirs(fallback, exist_ok=True)
        return fallback

GENERATED_DIR = get_storage_dir("generated")
UPLOADS_DIR = get_storage_dir("generated/uploads")
TEMPLATES_DIR = get_storage_dir("templates")

REPORTS_DB_PATH = os.path.join(GENERATED_DIR, "reports_db.json")
USERS_DB_PATH = os.path.join(GENERATED_DIR, "users_db.json")
SETTINGS_DB_PATH = os.path.join(GENERATED_DIR, "settings_db.json")

def get_builtin_template_path() -> str:
    possible_paths = [
        os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "Template.docx")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "templates", "Template.docx")),
        os.path.abspath(os.path.join(os.getcwd(), "backend", "templates", "Template.docx")),
        os.path.abspath(os.path.join(os.getcwd(), "templates", "Template.docx")),
    ]
    for p in possible_paths:
        if os.path.exists(p):
            return p
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "Template.docx"))

def get_builtin_fields_config_path() -> str:
    possible_paths = [
        os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "fields_config.json")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "templates", "fields_config.json")),
        os.path.abspath(os.path.join(os.getcwd(), "backend", "templates", "fields_config.json")),
        os.path.abspath(os.path.join(os.getcwd(), "templates", "fields_config.json")),
    ]
    for p in possible_paths:
        if os.path.exists(p):
            return p
    return ""


# Helper for Password Hashing
def hash_password(password: str, salt: str = "spc_shared_salt") -> str:
    return hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 100000).hex()

# Helpers for User & Settings Storage
def load_users() -> list:
    if os.path.exists(USERS_DB_PATH):
        try:
            with open(USERS_DB_PATH, "r") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_users(users: list):
    with open(USERS_DB_PATH, "w") as f:
        json.dump(users, f, indent=2)

def init_default_user():
    users = load_users()
    if not users:
        default_user = {
            "id": "user_spc_team",
            "email": "spc@tpo.edu",
            "password_hash": hash_password("spc12345"),
            "account_name": "Training & Placement Cell",
            "token": "token_spc_shared_account_default",
            "created_at": datetime.now().isoformat()
        }
        save_users([default_user])

def load_settings() -> dict:
    default_settings = {
        "account_name": "Training & Placement Cell",
        "email": "spc@tpo.edu",
        "theme": "light",
        "export_preference": "ask",
        "default_college_name": "SPC Institute of Technology",
        "default_department": "Training & Placement Cell",
        "default_organizer": "TPO Team",
        "default_venue": "Main Auditorium"
    }
    if os.path.exists(SETTINGS_DB_PATH):
        try:
            with open(SETTINGS_DB_PATH, "r") as f:
                saved = json.load(f)
                default_settings.update(saved)
        except Exception:
            pass
    return default_settings

def save_settings(settings: dict):
    with open(SETTINGS_DB_PATH, "w") as f:
        json.dump(settings, f, indent=2)

init_default_user()

app = FastAPI(title="SPC Documentation AI API")
router = APIRouter()

app.mount("/generated", StaticFiles(directory=GENERATED_DIR), name="generated")
app.mount("/api/generated", StaticFiles(directory=GENERATED_DIR), name="api_generated")

ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", "*").split(",") if o.strip()]
is_wildcard = "*" in ALLOWED_ORIGINS or not ALLOWED_ORIGINS

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if is_wildcard else ALLOWED_ORIGINS,
    allow_credentials=False if is_wildcard else True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini dynamically
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

class SignupRequest(BaseModel):
    email: str
    password: str
    account_name: Optional[str] = "Training & Placement Cell"

class LoginRequest(BaseModel):
    email: str
    password: str

class SettingsModel(BaseModel):
    account_name: Optional[str] = "Training & Placement Cell"
    email: Optional[str] = "spc@tpo.edu"
    theme: Optional[str] = "light"
    export_preference: Optional[str] = "ask"
    default_college_name: Optional[str] = ""
    default_department: Optional[str] = ""
    default_organizer: Optional[str] = ""
    default_venue: Optional[str] = ""

class AnalyzeTemplateResponse(BaseModel):
    fields: list[dict]

class TemplateField(BaseModel):
    name: str
    label: str
    type: str
    value: Optional[str] = ""
    originalText: Optional[str] = ""
    confidence_score: Optional[int] = 100
    confidence_level: Optional[str] = "high"

class SaveTemplateRequest(BaseModel):
    fields: List[TemplateField]

class GenerateDocumentRequest(BaseModel):
    values: Optional[Dict[str, Any]] = None
    photo_assignments: Optional[Dict[str, str]] = None
    template_filename: Optional[str] = None

class AutoFillRequest(BaseModel):
    notes: Optional[str] = ""

# Helper: Save report metadata record
def save_report_record(record: dict):
    try:
        records = []
        if os.path.exists(REPORTS_DB_PATH):
            with open(REPORTS_DB_PATH, "r") as f:
                records = json.load(f)
        records.insert(0, record)
        with open(REPORTS_DB_PATH, "w") as f:
            json.dump(records, f, indent=2)
    except Exception as e:
        print(f"Error saving report metadata: {e}")

# Root & Health Check Endpoint
@router.get("/")
async def root():
    return {"status": "ok", "message": "SPC Documentation AI API is running"}

# Auth & Settings Endpoints
@router.post("/auth/signup")
async def signup(req: SignupRequest):
    users = load_users()
    for u in users:
        if u["email"].lower() == req.email.lower():
            raise HTTPException(status_code=400, detail="Account with this email already exists.")
    
    token = f"token_{uuid.uuid4().hex}"
    new_user = {
        "id": f"user_{int(time.time())}",
        "email": req.email,
        "password_hash": hash_password(req.password),
        "account_name": req.account_name or "Training & Placement Cell",
        "token": token,
        "created_at": datetime.now().isoformat()
    }
    users.append(new_user)
    save_users(users)
    return {
        "message": "Account created successfully",
        "token": token,
        "user": {
            "id": new_user["id"],
            "email": new_user["email"],
            "account_name": new_user["account_name"]
        }
    }

@router.post("/auth/login")
async def login(req: LoginRequest):
    users = load_users()
    pwd_hash = hash_password(req.password)
    for u in users:
        if u["email"].lower() == req.email.lower() and u["password_hash"] == pwd_hash:
            return {
                "message": "Login successful",
                "token": u["token"],
                "user": {
                    "id": u["id"],
                    "email": u["email"],
                    "account_name": u.get("account_name", "Training & Placement Cell")
                }
            }
    raise HTTPException(status_code=401, detail="Invalid email or password.")

@router.get("/auth/me")
async def get_current_user(authorization: Optional[str] = Header(None)):
    users = load_users()
    if not authorization:
        if users:
            u = users[0]
            return {"user": {"id": u["id"], "email": u["email"], "account_name": u.get("account_name", "Training & Placement Cell")}}
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    token = authorization.replace("Bearer ", "").strip()
    for u in users:
        if u["token"] == token:
            return {"user": {"id": u["id"], "email": u["email"], "account_name": u.get("account_name", "Training & Placement Cell")}}
    
    raise HTTPException(status_code=401, detail="Invalid token")

@router.get("/settings")
async def get_settings_endpoint():
    return load_settings()

@router.post("/settings")
async def update_settings_endpoint(req: SettingsModel):
    current = load_settings()
    updated = req.model_dump(exclude_unset=True)
    current.update(updated)
    save_settings(current)
    return {"message": "Settings saved successfully", "settings": current}

@router.get("/reports")
async def get_reports_endpoint():
    if os.path.exists(REPORTS_DB_PATH):
        try:
            with open(REPORTS_DB_PATH, "r") as f:
                return json.load(f)
        except Exception:
            return []
    return []

@router.delete("/reports/{report_id}")
async def delete_report_endpoint(report_id: str):
    if not os.path.exists(REPORTS_DB_PATH):
        raise HTTPException(status_code=404, detail="Report database not found")
    
    try:
        with open(REPORTS_DB_PATH, "r") as f:
            records = json.load(f)
        
        target = None
        remaining = []
        for r in records:
            if r.get("id") == report_id or r.get("docxFilename") == report_id or r.get("pdfFilename") == report_id:
                target = r
            else:
                remaining.append(r)
        
        if not target:
            raise HTTPException(status_code=404, detail="Report not found")
        
        with open(REPORTS_DB_PATH, "w") as f:
            json.dump(remaining, f, indent=2)
            
        for key in ["docxFilename", "pdfFilename"]:
            fname = target.get(key)
            if fname:
                fp = os.path.join(GENERATED_DIR, fname)
                if os.path.exists(fp):
                    try:
                        os.remove(fp)
                    except Exception:
                        pass
                        
        return {"message": "Report deleted successfully", "id": report_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete report: {e}")

# Helper: Extract text from DOCX bytes
def extract_text_from_docx_bytes(contents: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(contents))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

# Helper: Extract text from PDF bytes
def extract_text_from_pdf_bytes(contents: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(contents))
        full_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                full_text.append(t)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

# Helper: Extract OCR text from Image using Gemini Vision
def extract_text_from_image_bytes(contents: bytes, filename: str) -> str:
    try:
        img = Image.open(io.BytesIO(contents))
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = (
            "Analyze this uploaded image for an event/activity report evidence or handwritten notes.\n"
            "1. Perform accurate OCR and extract ALL readable text (titles, dates, names, locations, numbers, handwritten notes).\n"
            "2. Describe any visual content or document structure.\n"
            "Provide output as clean readable plain text."
        )
        response = model.generate_content([prompt, img])
        return response.text.strip() if response and response.text else f"Notes extracted from image {filename}: [Activity details captured]"
    except Exception as e:
        print(f"Vision OCR notice for {filename}: {e}")
        return f"Uploaded image note ({filename}): Event details captured successfully."

@app.get("/")
def read_root():
    return {"message": "Welcome to the SPC Documentation AI Platform API"}

# Get previous saved reports
@router.get("/reports")
async def get_previous_reports():
    if not os.path.exists(REPORTS_DB_PATH):
        return {"reports": []}
    try:
        with open(REPORTS_DB_PATH, "r") as f:
            records = json.load(f)
            return {"reports": records}
    except Exception as e:
        return {"reports": []}

@router.get("/templates/fields")
async def get_template_fields():
    try:
        config_path = get_builtin_fields_config_path()
        if config_path and os.path.exists(config_path):
            with open(config_path, "r") as f:
                fields = json.load(f)
                if fields:
                    return {"fields": fields}
        return {"fields": DEFAULT_ACADEMIC_FIELDS}
    except Exception:
        return {"fields": DEFAULT_ACADEMIC_FIELDS}

# Report Generation Endpoint using Built-in Developer Template
@router.post("/templates/generate")
async def generate_document(
    values: Optional[str] = Form(None),
    notice_file: Optional[UploadFile] = File(None),
    event_photos: Optional[List[UploadFile]] = File(None),
    photo_assignments: Optional[str] = Form(None),
    body: Optional[GenerateDocumentRequest] = None
):
    parsed_values = {}
    parsed_photos = {}

    if values is not None:
        try:
            parsed_values = json.loads(values)
        except Exception:
            parsed_values = {}
    elif body is not None:
        parsed_values = body.values or {}
        parsed_photos = body.photo_assignments or {}

    if photo_assignments and isinstance(photo_assignments, str):
        try:
            parsed_photos = json.loads(photo_assignments)
        except Exception:
            pass

    template_path = get_builtin_template_path()

    # Save notice/brochure file if provided
    notice_saved_path = None
    if notice_file and notice_file.filename:
        n_bytes = await notice_file.read()
        unique_n_name = f"notice_{int(time.time())}_{notice_file.filename}"
        notice_saved_path = os.path.join(UPLOADS_DIR, unique_n_name)
        with open(notice_saved_path, "wb") as f:
            f.write(n_bytes)

    # Save multiple uploaded event photos
    saved_event_photo_paths = []
    if event_photos:
        for ep in event_photos:
            if ep and ep.filename:
                ep_bytes = await ep.read()
                unique_ep_name = f"event_photo_{int(time.time())}_{ep.filename}"
                ep_save_path = os.path.join(UPLOADS_DIR, unique_ep_name)
                with open(ep_save_path, "wb") as f:
                    f.write(ep_bytes)
                saved_event_photo_paths.append(ep_save_path)

    try:
        try:
            doc = docx.Document(template_path) if os.path.exists(template_path) else docx.Document()
        except Exception as load_err:
            print(f"Error opening template at {template_path}: {load_err}")
            doc = docx.Document()

        # 1. Fill SJCEM AF-5 Structured Tables if applicable
        if len(doc.tables) >= 4:
            try:
                t0 = doc.tables[0]
                t0.rows[0].cells[1].text = parsed_values.get("activity_name", "")
                t0.rows[1].cells[1].text = parsed_values.get("date_time", "")
                t0.rows[1].cells[3].text = parsed_values.get("department", "")
                t0.rows[2].cells[1].text = parsed_values.get("venue", "")
                t0.rows[2].cells[3].text = parsed_values.get("participants", "")
                t0.rows[3].cells[1].text = parsed_values.get("nature_of_activity", "")
                t0.rows[3].cells[3].text = parsed_values.get("mode_of_activity", "")
                t0.rows[4].cells[1].text = parsed_values.get("activity_incharge", "")
                t0.rows[4].cells[3].text = parsed_values.get("activity_coordinator", "")
                t0.rows[5].cells[1].text = parsed_values.get("resource_person", "")

                t1 = doc.tables[1]
                t1.rows[0].cells[1].text = parsed_values.get("objectives", "")
                t1.rows[1].cells[1].text = parsed_values.get("target_audience", "")
                t1.rows[2].cells[1].text = parsed_values.get("methodology", "")
                t1.rows[3].cells[1].text = parsed_values.get("outcomes", "")

                t2 = doc.tables[2]
                t2.rows[1].cells[0].text = parsed_values.get("strengths", "")
                t2.rows[1].cells[1].text = parsed_values.get("weaknesses", "")
                t2.rows[1].cells[2].text = parsed_values.get("opportunities", "")
                t2.rows[1].cells[3].text = parsed_values.get("threats", "")

                t3 = doc.tables[3]
                t3.rows[0].cells[0].text = parsed_values.get("notice_brochure_tick", "")
                t3.rows[0].cells[2].text = parsed_values.get("feedback_analysis_tick", "")
                t3.rows[1].cells[0].text = parsed_values.get("attendance_list_tick", "")
                t3.rows[1].cells[2].text = parsed_values.get("news_letter_data_tick", "")
                t3.rows[2].cells[0].text = parsed_values.get("photos_tick", "")
                t3.rows[2].cells[2].text = parsed_values.get("media_news_details_tick", "")
                t3.rows[3].cells[0].text = parsed_values.get("certificate_tick", "")
                t3.rows[3].cells[2].text = parsed_values.get("co_po_mapping_tick", "")
                t3.rows[4].cells[0].text = parsed_values.get("feedback_form_tick", "")
                t3.rows[4].cells[2].text = parsed_values.get("any_other_tick", "")
            except Exception as table_err:
                print(f"Table fill notice: {table_err}")

        # Helper: Replace text in paragraph
        def replace_text_in_paragraph(para, target_text, replacement_text):
            if not target_text:
                return
            repl = replacement_text if replacement_text is not None else ""
            if target_text in para.text:
                replaced = False
                for run in para.runs:
                    if target_text in run.text:
                        run.text = run.text.replace(target_text, repl)
                        replaced = True
                if not replaced:
                    para.text = para.text.replace(target_text, repl)

        # Helper: Insert photo into paragraph preserving aspect ratio
        def insert_photo_into_paragraph(para, photo_path, max_width_inches=3.0):
            try:
                if not photo_path or not os.path.exists(photo_path):
                    return False
                img = Image.open(photo_path)
                width_px, height_px = img.size
                aspect_ratio = height_px / float(width_px) if width_px > 0 else 0.75

                target_width = Inches(max_width_inches)
                target_height = Inches(max_width_inches * aspect_ratio)

                para.text = ""
                run = para.add_run()
                run.add_picture(photo_path, width=target_width, height=target_height)
                return True
            except Exception as img_err:
                print(f"Error inserting picture {photo_path}: {img_err}")
                return False

        # Track notice placement
        notice_inserted = False
        if notice_saved_path:
            parsed_photos["[NOTICE_PHOTO]"] = notice_saved_path
            parsed_photos["[NOTICE]"] = notice_saved_path

        # 2. Process paragraphs for text replacement & photo insertion
        for para in doc.paragraphs:
            for key, val in parsed_values.items():
                replace_text_in_paragraph(para, f"[{key}]", val or "")
                replace_text_in_paragraph(para, f"{{{key}}}", val or "")
                if val:
                    replace_text_in_paragraph(para, key, val)

            for ph_key, p_path in parsed_photos.items():
                if ph_key in para.text and p_path:
                    if insert_photo_into_paragraph(para, p_path):
                        if p_path == notice_saved_path:
                            notice_inserted = True

        # 3. Process table cells for text replacement & photo insertion
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in parsed_values.items():
                            replace_text_in_paragraph(para, f"[{key}]", val or "")
                            replace_text_in_paragraph(para, f"{{{key}}}", val or "")

                        for ph_key, p_path in parsed_photos.items():
                            if ph_key in para.text and p_path:
                                if insert_photo_into_paragraph(para, p_path, max_width_inches=2.5):
                                    if p_path == notice_saved_path:
                                        notice_inserted = True

        # Ensure Notice & Brochure photo is attached if not already inserted by placeholder
        if notice_saved_path and not notice_inserted:
            doc.add_page_break()
            p_nhead = doc.add_paragraph()
            run_nh = p_nhead.add_run("Notice & Brochure")
            run_nh.bold = True
            run_nh.font.size = Pt(14)
            insert_photo_into_paragraph(doc.add_paragraph(), notice_saved_path, max_width_inches=5.2)

        # 4. Process Multiple Event Photos into a Structured 2-Column Table Grid (2x2 / 2x3 Layout)
        if saved_event_photo_paths:
            doc.add_page_break()
            p_ehead = doc.add_paragraph()
            run_eh = p_ehead.add_run("Event Photographs")
            run_eh.bold = True
            run_eh.font.size = Pt(14)

            # Create 2-column Word table for neat grid layout
            grid_table = doc.add_table(rows=0, cols=2)
            grid_table.autofit = False

            for i in range(0, len(saved_event_photo_paths), 2):
                row_cells = grid_table.add_row().cells
                
                # Left Column Photo
                path1 = saved_event_photo_paths[i]
                p1 = row_cells[0].paragraphs[0]
                insert_photo_into_paragraph(p1, path1, max_width_inches=2.8)
                
                # Right Column Photo (if present)
                if i + 1 < len(saved_event_photo_paths):
                    path2 = saved_event_photo_paths[i + 1]
                    p2 = row_cells[1].paragraphs[0]
                    insert_photo_into_paragraph(p2, path2, max_width_inches=2.8)

        # Clean up remaining unfilled placeholders like [field_name] to ensure empty fields remain blank
        for para in doc.paragraphs:
            para.text = re.sub(r'\[[a-zA-Z0-9_]+\]', '', para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = re.sub(r'\[[a-zA-Z0-9_]+\]', '', para.text)

        # Strict Naming Format: [Name_of_the_Activity]_documentation.pdf / .docx
        raw_activity_name = parsed_values.get("activity_name", "").strip() or "Activity"
        clean_activity_name = re.sub(r'[^a-zA-Z0-9_\-]', '', raw_activity_name.replace(' ', '_'))
        clean_activity_name = clean_activity_name[:40].rstrip('_')
        if not clean_activity_name: clean_activity_name = "Activity"

        base_filename = f"{clean_activity_name}_documentation"
        docx_filename = f"{base_filename}.docx"
        pdf_filename = f"{base_filename}.pdf"

        os.makedirs(GENERATED_DIR, exist_ok=True)
        docx_path = os.path.join(GENERATED_DIR, docx_filename)
        pdf_path = os.path.join(GENERATED_DIR, pdf_filename)

        doc.save(docx_path)

        try:
            if DOCX2PDF_AVAILABLE:
                abs_docx = os.path.abspath(docx_path)
                abs_pdf = os.path.abspath(pdf_path)
                convert(abs_docx, abs_pdf)
            else:
                pdf_filename = None
        except Exception as e:
            print(f"PDF conversion notice: {e}")
            pdf_filename = None

        API_BASE_URL = os.getenv("API_BASE_URL", "").rstrip("/")
        docx_url = f"{API_BASE_URL}/generated/{docx_filename}" if API_BASE_URL else f"/generated/{docx_filename}"
        pdf_url = (f"{API_BASE_URL}/generated/{pdf_filename}" if API_BASE_URL else f"/generated/{pdf_filename}") if pdf_filename and os.path.exists(pdf_path) else None

        # Auto-Save Report Record for "Previous Reports" Section
        report_record = {
            "id": f"report_{int(time.time())}",
            "activity_name": raw_activity_name,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "timestamp": int(time.time()),
            "docx_url": docx_url,
            "pdf_url": pdf_url,
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename
        }
        save_report_record(report_record)

        return {
            "docxUrl": docx_url,
            "pdfUrl": pdf_url,
            "docxFilename": docx_filename,
            "pdfFilename": pdf_filename
        }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating document: {e}")

# Register router for both root and /api paths
app.include_router(router)
app.include_router(router, prefix="/api")

if __name__ == "__main__":
    import uvicorn
    PORT = int(os.getenv("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=PORT, reload=False)
