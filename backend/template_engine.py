import os
import json
import re
import traceback
from typing import Dict, Any, List, Optional
from PIL import Image
import docx
from docx.shared import Inches, Pt

def get_default_mapping_path() -> str:
    possible_paths = [
        os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "template_mapping.json")),
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "templates", "template_mapping.json")),
        os.path.abspath(os.path.join(os.getcwd(), "backend", "templates", "template_mapping.json")),
        os.path.abspath(os.path.join(os.getcwd(), "templates", "template_mapping.json")),
    ]
    for p in possible_paths:
        if os.path.exists(p):
            return p
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "template_mapping.json"))

def load_template_mapping(config_path: Optional[str] = None) -> Dict[str, Any]:
    target_path = config_path or get_default_mapping_path()
    if target_path and os.path.exists(target_path):
        try:
            with open(target_path, "r") as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading template_mapping.json from {target_path}: {e}")
    return {}

def insert_photo_into_paragraph(para, photo_path: str, max_width_inches: float = 3.0) -> bool:
    """
    Inserts a photo into a docx paragraph while preserving exact image aspect ratio.
    """
    try:
        if not photo_path or not os.path.exists(photo_path):
            return False
        img = Image.open(photo_path)
        width_px, height_px = img.size
        if width_px <= 0 or height_px <= 0:
            return False
            
        aspect_ratio = height_px / float(width_px)
        target_width = Inches(max_width_inches)
        target_height = Inches(max_width_inches * aspect_ratio)

        para.text = ""
        run = para.add_run()
        run.add_picture(photo_path, width=target_width, height=target_height)
        return True
    except Exception as img_err:
        print(f"Error inserting picture {photo_path}: {img_err}")
        return False

def replace_text_in_paragraph(para, target_text: str, replacement_text: str):
    if not target_text:
        return
    repl = replacement_text if replacement_text is not None else ""
    if target_text in para.text:
        replaced = False
        for run in para.runs:
            if target_text in run.text:
                run.text = run.text.replace(target_text, repl)
                replaced = True
        if not replaced:
            para.text = para.text.replace(target_text, repl)

def generate_docx_report(
    template_path: str,
    output_docx_path: str,
    parsed_values: Dict[str, Any],
    notice_photo_path: Optional[str] = None,
    event_photo_paths: Optional[List[str]] = None,
    feedback_graph_path: Optional[str] = None,
    feedback_interpretation: Optional[str] = None,
    mapping_config_path: Optional[str] = None
) -> str:
    """
    Generates a DOCX report by populating template tables and image placements
    driven entirely by configuration (template_mapping.json).
    """
    try:
        if os.path.exists(template_path):
            doc = docx.Document(template_path)
        else:
            print(f"Template path {template_path} not found. Creating empty Document.")
            doc = docx.Document()

        mapping = load_template_mapping(mapping_config_path)
        field_mappings = mapping.get("field_mappings", [])
        image_placements = mapping.get("image_placements", {})

        # 1. Configuration-driven Table Cell Population
        for fm in field_mappings:
            if fm.get("type") == "table_cell":
                tbl_idx = fm.get("table_index", 0)
                row_idx = fm.get("row", 0)
                col_idx = fm.get("col", 0)
                field_name = fm.get("field", "")

                val = str(parsed_values.get(field_name, ""))
                
                if tbl_idx < len(doc.tables):
                    tbl = doc.tables[tbl_idx]
                    if row_idx < len(tbl.rows):
                        row = tbl.rows[row_idx]
                        if col_idx < len(row.cells):
                            row.cells[col_idx].text = val

        # 2. Configuration-driven Paragraph & Cell Placeholder Replacement
        formats = mapping.get("placeholder_formats", ["[{key}]", "{{{key}}}"])
        for para in doc.paragraphs:
            for key, val in parsed_values.items():
                if val is not None:
                    for fmt in formats:
                        target = fmt.replace("{key}", key)
                        replace_text_in_paragraph(para, target, str(val))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in parsed_values.items():
                            if val is not None:
                                for fmt in formats:
                                    target = fmt.replace("{key}", key)
                                    replace_text_in_paragraph(para, target, str(val))

        # 3. Configuration-driven Notice & Brochure Attachment
        notice_cfg = image_placements.get("notice_brochure", {})
        notice_placeholders = notice_cfg.get("placeholders", ["[NOTICE_PHOTO]", "[NOTICE]"])
        notice_max_w = notice_cfg.get("max_width_inches", 5.2)
        notice_heading = notice_cfg.get("fallback_heading", "Notice & Brochure")

        notice_inserted = False
        if notice_photo_path and os.path.exists(notice_photo_path):
            for para in doc.paragraphs:
                if any(ph in para.text for ph in notice_placeholders) and not notice_inserted:
                    notice_inserted = insert_photo_into_paragraph(para, notice_photo_path, max_width_inches=notice_max_w)

            if not notice_inserted:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if any(ph in para.text for ph in notice_placeholders) and not notice_inserted:
                                    notice_inserted = insert_photo_into_paragraph(para, notice_photo_path, max_width_inches=4.5)

            if not notice_inserted:
                doc.add_page_break()
                p_nhead = doc.add_paragraph()
                run_nh = p_nhead.add_run(notice_heading)
                run_nh.bold = True
                run_nh.font.size = Pt(14)
                insert_photo_into_paragraph(doc.add_paragraph(), notice_photo_path, max_width_inches=notice_max_w)

        # 4. Configuration-driven Feedback Analysis Section
        fb_cfg = image_placements.get("feedback_graph", {})
        fb_placeholders = fb_cfg.get("placeholders", ["[FEEDBACK_GRAPH]", "[FEEDBACK_ANALYSIS]"])
        fb_max_w = fb_cfg.get("max_width_inches", 5.0)
        fb_heading = fb_cfg.get("fallback_heading", "Feedback Analysis")

        feedback_inserted = False
        interp_text = feedback_interpretation or parsed_values.get("feedback_interpretation") or parsed_values.get("feedback_summary") or ""

        if feedback_graph_path and os.path.exists(feedback_graph_path):
            for para in doc.paragraphs:
                if any(ph in para.text for ph in fb_placeholders) and not feedback_inserted:
                    feedback_inserted = insert_photo_into_paragraph(para, feedback_graph_path, max_width_inches=fb_max_w)
                    if interp_text:
                        p_interp = doc.add_paragraph()
                        r_interp = p_interp.add_run(f"Summary & Interpretation:\n{interp_text}")
                        r_interp.font.size = Pt(10.5)

            if not feedback_inserted:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if any(ph in para.text for ph in fb_placeholders) and not feedback_inserted:
                                    feedback_inserted = insert_photo_into_paragraph(para, feedback_graph_path, max_width_inches=4.5)
                                    if interp_text:
                                        p_interp = cell.add_paragraph()
                                        r_interp = p_interp.add_run(f"Summary & Interpretation:\n{interp_text}")
                                        r_interp.font.size = Pt(10.5)

            if not feedback_inserted:
                doc.add_page_break()
                p_fhead = doc.add_paragraph()
                run_fh = p_fhead.add_run(fb_heading)
                run_fh.bold = True
                run_fh.font.size = Pt(14)

                insert_photo_into_paragraph(doc.add_paragraph(), feedback_graph_path, max_width_inches=fb_max_w)

                if interp_text:
                    p_desc = doc.add_paragraph()
                    run_desc = p_desc.add_run(interp_text)
                    run_desc.font.size = Pt(10.5)
        elif interp_text:
            doc.add_page_break()
            p_fhead = doc.add_paragraph()
            run_fh = p_fhead.add_run(fb_heading)
            run_fh.bold = True
            run_fh.font.size = Pt(14)

            p_desc = doc.add_paragraph()
            run_desc = p_desc.add_run(interp_text)
            run_desc.font.size = Pt(10.5)

        # 5. Configuration-driven Event Photographs Section
        ep_cfg = image_placements.get("event_photos", {})
        ep_placeholders = ep_cfg.get("placeholders", ["[EVENT_PHOTOS]"])
        ep_max_w = ep_cfg.get("max_width_inches", 2.8)
        ep_heading = ep_cfg.get("fallback_heading", "Event Photographs")

        if event_photo_paths:
            doc.add_page_break()
            p_ehead = doc.add_paragraph()
            run_eh = p_ehead.add_run(ep_heading)
            run_eh.bold = True
            run_eh.font.size = Pt(14)

            grid_table = doc.add_table(rows=0, cols=2)
            grid_table.autofit = False

            for i in range(0, len(event_photo_paths), 2):
                row_cells = grid_table.add_row().cells
                
                path1 = event_photo_paths[i]
                p1 = row_cells[0].paragraphs[0]
                insert_photo_into_paragraph(p1, path1, max_width_inches=ep_max_w)
                
                if i + 1 < len(event_photo_paths):
                    path2 = event_photo_paths[i + 1]
                    p2 = row_cells[1].paragraphs[0]
                    insert_photo_into_paragraph(p2, path2, max_width_inches=ep_max_w)

        # 6. Clean up remaining unfilled placeholders like [field_name] to ensure empty fields remain blank
        for para in doc.paragraphs:
            para.text = re.sub(r'\[[a-zA-Z0-9_]+\]', '', para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = re.sub(r'\[[a-zA-Z0-9_]+\]', '', para.text)

        os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
        doc.save(output_docx_path)
        return output_docx_path

    except Exception as e:
        print(f"Error generating DOCX report: {e}")
        traceback.print_exc()
        raise e
